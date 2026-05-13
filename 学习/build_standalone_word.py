from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "2026-05-13-前四轮总结-甲方视角介绍与批判.md"
EXPORT_DIR = ROOT / "独立导出"
GENERATED_DIR = EXPORT_DIR / "generated"
TEMP_MD = EXPORT_DIR / "TradingAgents-前四轮总结-独立版.md"
OUTPUT_DOCX = EXPORT_DIR / "TradingAgents-前四轮总结-独立版.docx"


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/msyhbd.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size)
    return ImageFont.load_default()


def draw_centered_text(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str) -> None:
    lines = text.split("\n")
    line_heights = []
    line_widths = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_widths.append(bbox[2] - bbox[0])
        line_heights.append(bbox[3] - bbox[1])

    total_height = sum(line_heights) + max(0, len(lines) - 1) * 6
    box_x1, box_y1, box_x2, box_y2 = box
    current_y = box_y1 + (box_y2 - box_y1 - total_height) // 2

    for line, line_width, line_height in zip(lines, line_widths, line_heights):
        current_x = box_x1 + (box_x2 - box_x1 - line_width) // 2
        draw.text((current_x, current_y), line, font=font, fill=fill)
        current_y += line_height + 6


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str, width: int = 4) -> None:
    draw.line([start, end], fill=color, width=width)
    arrow_size = 10
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex >= sx else -1
        points = [(ex, ey), (ex - direction * arrow_size, ey - arrow_size), (ex - direction * arrow_size, ey + arrow_size)]
    else:
        direction = 1 if ey >= sy else -1
        points = [(ex, ey), (ex - arrow_size, ey - direction * arrow_size), (ex + arrow_size, ey - direction * arrow_size)]
    draw.polygon(points, fill=color)


def rounded_box(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, font, fill: str, outline: str) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    draw_centered_text(draw, box, text, font, fill="#1f2937")


def build_horizontal_flowchart(output_path: Path) -> None:
    image = Image.new("RGB", (2200, 520), "white")
    draw = ImageDraw.Draw(image)
    font = load_font(28)
    label_font = load_font(38)

    boxes = [
        (40, 170, 270, 340, "输入 ticker\n与日期"),
        (320, 140, 630, 370, "Analyst Team\n技术面 情绪面\n新闻面 基本面"),
        (690, 140, 980, 370, "Bull / Bear\nResearcher\n多空辩论"),
        (1040, 140, 1320, 370, "Research Manager\n形成投资计划"),
        (1380, 140, 1600, 370, "Trader\n生成交易提案"),
        (1660, 120, 1950, 390, "Risk Team\n激进 保守\n中性辩论"),
        (2010, 140, 2160, 370, "Portfolio\nManager"),
    ]

    palette = ["#dbeafe", "#dcfce7", "#fee2e2", "#ede9fe", "#fef3c7", "#cffafe", "#fde68a"]
    for (x1, y1, x2, y2, text), fill in zip(boxes, palette):
        rounded_box(draw, (x1, y1, x2, y2), text, font, fill, "#94a3b8")

    for current, nxt in zip(boxes, boxes[1:]):
        draw_arrow(draw, (current[2] + 12, (current[1] + current[3]) // 2), (nxt[0] - 12, (nxt[1] + nxt[3]) // 2), "#475569")

    draw.text((40, 40), "TradingAgents 主流程图", font=label_font, fill="#111827")
    image.save(output_path)


def build_vertical_flowchart(output_path: Path) -> None:
    image = Image.new("RGB", (1200, 1600), "white")
    draw = ImageDraw.Draw(image)
    font = load_font(30)
    label_font = load_font(42)

    nodes = [
        (330, 120, 870, 240, "开始分析"),
        (220, 310, 980, 460, "四类分析师产出报告"),
        (260, 530, 940, 680, "多空研究员辩论"),
        (260, 750, 940, 900, "研究经理输出计划"),
        (280, 970, 920, 1120, "交易员提出方案"),
        (240, 1190, 960, 1340, "三类风险角色讨论"),
        (260, 1410, 940, 1560, "组合经理最终裁决\n并写入日志与反思记忆"),
    ]

    palette = ["#e5e7eb", "#dbeafe", "#dcfce7", "#ede9fe", "#fef3c7", "#cffafe", "#fee2e2"]
    for (x1, y1, x2, y2, text), fill in zip(nodes, palette):
        rounded_box(draw, (x1, y1, x2, y2), text, font, fill, "#94a3b8")

    for current, nxt in zip(nodes, nodes[1:]):
        draw_arrow(draw, ((current[0] + current[2]) // 2, current[3] + 12), ((nxt[0] + nxt[2]) // 2, nxt[1] - 12), "#475569")

    draw.text((40, 40), "TradingAgents 简化角色流程图", font=label_font, fill="#111827")
    image.save(output_path)


def build_standalone_markdown(source_text: str) -> str:
    main_mermaid = """```mermaid
flowchart LR
\tA[输入 ticker 与日期] --> B[Analyst Team
技术面 情绪面 新闻面 基本面]
\tB --> C[Bull / Bear Researcher
多空辩论]
\tC --> D[Research Manager
形成投资计划]
\tD --> E[Trader
生成交易提案]
\tE --> F[Risk Team
激进 保守 中性辩论]
\tF --> G[Portfolio Manager
输出最终决策]
\tG --> H[日志 记忆 反思]
```"""

    role_mermaid = """```mermaid
flowchart TD
\tStart[开始分析] --> Analyst[四类分析师产出报告]
\tAnalyst --> Debate[多空研究员辩论]
\tDebate --> Plan[研究经理输出计划]
\tPlan --> Trade[交易员提出方案]
\tTrade --> Risk[三类风险角色讨论]
\tRisk --> PM[组合经理最终裁决]
\tPM --> Memory[写入日志与反思记忆]
```"""

    transformed = source_text.replace(main_mermaid, "![TradingAgents 主流程图](generated/main_flowchart.png)")
    transformed = transformed.replace(role_mermaid, "![TradingAgents 简化角色流程图](generated/role_flowchart.png)")
    transformed = transformed.replace("../源代码/assets/", "../../源代码/assets/")
    return transformed


def set_document_defaults(document: Document) -> None:
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style.font.size = Pt(11)
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_page_break_rule(paragraph) -> None:
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_after = Pt(10)
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    border.append(bottom)
    paragraph._p.get_or_add_pPr().append(border)


def clean_inline_markdown(text: str) -> str:
    text = re.sub(r"!\[[^\]]*\]\(([^)]+)\)", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


def add_image(document: Document, image_path: Path) -> None:
    with Image.open(image_path) as image:
        width_px, height_px = image.size

    max_width_inches = 6.5
    width_inches = min(max_width_inches, width_px / 160)
    document.add_picture(str(image_path), width=Inches(width_inches))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def flush_paragraph(document: Document, lines: list[str]) -> None:
    if not lines:
        return
    text = clean_inline_markdown(" ".join(line.strip() for line in lines))
    if text:
        document.add_paragraph(text)
    lines.clear()


def export_docx(markdown_text: str, output_path: Path) -> None:
    document = Document()
    set_document_defaults(document)

    paragraph_buffer: list[str] = []
    in_code_block = False

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            continue

        if not stripped:
            flush_paragraph(document, paragraph_buffer)
            continue

        if stripped == "---":
            flush_paragraph(document, paragraph_buffer)
            divider = document.add_paragraph()
            add_page_break_rule(divider)
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            flush_paragraph(document, paragraph_buffer)
            level = min(len(heading_match.group(1)), 4)
            document.add_heading(clean_inline_markdown(heading_match.group(2)), level=level)
            continue

        image_match = re.match(r"^!\[[^\]]*\]\(([^)]+)\)$", stripped)
        if image_match:
            flush_paragraph(document, paragraph_buffer)
            image_path = (EXPORT_DIR / image_match.group(1)).resolve()
            add_image(document, image_path)
            continue

        ordered_match = re.match(r"^\d+\.\s+(.*)$", stripped)
        if ordered_match:
            flush_paragraph(document, paragraph_buffer)
            document.add_paragraph(clean_inline_markdown(ordered_match.group(1)), style="List Number")
            continue

        bullet_match = re.match(r"^-\s+(.*)$", stripped)
        if bullet_match:
            flush_paragraph(document, paragraph_buffer)
            document.add_paragraph(clean_inline_markdown(bullet_match.group(1)), style="List Bullet")
            continue

        paragraph_buffer.append(stripped)

    flush_paragraph(document, paragraph_buffer)
    document.save(output_path)


def main() -> None:
    EXPORT_DIR.mkdir(exist_ok=True)
    GENERATED_DIR.mkdir(exist_ok=True)

    build_horizontal_flowchart(GENERATED_DIR / "main_flowchart.png")
    build_vertical_flowchart(GENERATED_DIR / "role_flowchart.png")

    source_text = SOURCE_MD.read_text(encoding="utf-8")
    standalone_text = build_standalone_markdown(source_text)
    TEMP_MD.write_text(standalone_text, encoding="utf-8")

    export_docx(standalone_text, OUTPUT_DOCX)
    print(f"Created: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()