"""将第八轮学习-从零配置与运行指南.md 导出为 Word 文档"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "2026-05-13-05-00-00-第八轮学习-从零配置与运行指南.md"
OUTPUT_DOCX = ROOT / "独立导出" / "TradingAgents-从零配置与运行指南.docx"


# ──────────────────────────────────────────────
# 样式工具
# ──────────────────────────────────────────────

def set_document_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(11)
    try:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    except Exception:
        pass


def _set_cell_font(cell, bold=False, size=10):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(size)
            run.font.bold = bold


def _shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def add_divider(doc: Document):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ──────────────────────────────────────────────
# 行内 Markdown 解析 → runs
# ──────────────────────────────────────────────

def add_inline(paragraph, text: str):
    """解析行内 **bold** 和 `code`，写入 runs"""
    pattern = re.compile(r"(\*\*(.+?)\*\*|`([^`]+)`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            paragraph.add_run(text[last:m.start()])
        if m.group(2):  # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        else:  # code
            run = paragraph.add_run(m.group(3))
            run.font.name = "Courier New"
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        last = m.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def plain_text(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    return text.strip()


# ──────────────────────────────────────────────
# 表格渲染
# ──────────────────────────────────────────────

def render_table(doc: Document, rows: list[list[str]]):
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx >= n_cols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline(para, cell_text.strip())
            _set_cell_font(cell, bold=(r_idx == 0), size=10)
            if r_idx == 0:
                _shade_cell(cell, "D9E2F3")
    doc.add_paragraph()  # 表后空行


# ──────────────────────────────────────────────
# 代码块渲染
# ──────────────────────────────────────────────

def render_code_block(doc: Document, lines: list[str], lang: str):
    for line in lines:
        p = doc.add_paragraph(line)
        p.style = "Normal"
        run = p.runs[0] if p.runs else p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x17, 0x3A, 0x6B)
        # 浅灰背景
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F0F0F0")
        pPr.append(shd)
    doc.add_paragraph()


# ──────────────────────────────────────────────
# 主渲染逻辑
# ──────────────────────────────────────────────

def is_table_sep(row: list[str]) -> bool:
    return all(re.fullmatch(r":?-+:?", c.strip()) for c in row if c.strip())


def export_docx(md_text: str, output_path: Path):
    doc = Document()
    set_document_defaults(doc)

    lines = md_text.splitlines()
    i = 0
    table_rows: list[list[str]] = []
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    def flush_table():
        nonlocal table_rows
        if table_rows:
            render_table(doc, table_rows)
            table_rows = []

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── 代码块 ──
        if stripped.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                in_code = False
                render_code_block(doc, code_lines, code_lang)
            i += 1
            continue

        if in_code:
            code_lines.append(raw)
            i += 1
            continue

        # ── 空行 ──
        if not stripped:
            flush_table()
            i += 1
            continue

        # ── 分隔线 ──
        if stripped == "---":
            flush_table()
            add_divider(doc)
            i += 1
            continue

        # ── 标题 ──
        heading_m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_m:
            flush_table()
            level = min(len(heading_m.group(1)), 4)
            doc.add_heading(plain_text(heading_m.group(2)), level=level)
            i += 1
            continue

        # ── 表格行 ──
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c for c in stripped.strip("|").split("|")]
            if not is_table_sep(cells):
                table_rows.append(cells)
            i += 1
            continue
        else:
            flush_table()

        # ── 有序列表 ──
        ol_m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if ol_m:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, ol_m.group(1))
            i += 1
            continue

        # ── 无序列表（多级缩进）──
        ul_m = re.match(r"^(\s*)[-*]\s+(.+)$", raw)
        if ul_m:
            indent = len(ul_m.group(1))
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_inline(p, ul_m.group(2))
            i += 1
            continue

        # ── 普通段落 ──
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1

    flush_table()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"已生成：{output_path}")


def main():
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    export_docx(md_text, OUTPUT_DOCX)


if __name__ == "__main__":
    main()
