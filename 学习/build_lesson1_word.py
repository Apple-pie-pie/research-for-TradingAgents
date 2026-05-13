"""把第一轮学习文档导出为 Word，正确处理代码块和表格。"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "2026-05-13-00-01-55-第一轮学习-程序逻辑讲解.md"
EXPORT_DIR = ROOT / "独立导出"
OUTPUT_DOCX = EXPORT_DIR / "TradingAgents-第一轮学习-程序结构精讲.docx"


def set_document_defaults(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def clean_inline(text: str) -> str:
    """去掉 **bold**、`code`、[link](url) 的 markdown 语法，保留纯文本。"""
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text.strip()


def add_code_block(doc: Document, lines: list[str]) -> None:
    """把代码块内容以等宽字体段落写入 Word。"""
    joined = "\n".join(lines)
    para = doc.add_paragraph()
    run = para.add_run(joined)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x6E)
    # 灰色底纹
    from docx.oxml import OxmlElement
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F4F6")
    pPr.append(shd)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)


def add_table_from_md(doc: Document, rows: list[str]) -> None:
    """把 Markdown 表格（含分隔行）转换为 Word 表格。"""
    data_rows = [r for r in rows if not re.match(r"^\|[-: |]+\|$", r.strip())]
    if not data_rows:
        return

    parsed = []
    for row in data_rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        parsed.append(cells)

    max_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(parsed):
        for c_idx, cell_text in enumerate(row):
            if c_idx < max_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = clean_inline(cell_text)
                if r_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

    doc.add_paragraph()


def export_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    set_document_defaults(doc)

    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    in_table = False
    table_rows: list[str] = []
    para_buffer: list[str] = []

    def flush_para():
        if not para_buffer:
            return
        text = clean_inline(" ".join(para_buffer))
        if text:
            doc.add_paragraph(text)
        para_buffer.clear()

    def flush_table():
        nonlocal in_table
        if table_rows:
            add_table_from_md(doc, table_rows)
        table_rows.clear()
        in_table = False

    for raw in markdown_text.splitlines():
        line = raw.rstrip()
        stripped = line.strip()

        # ---------- 代码块 ----------
        if stripped.startswith("```"):
            if not in_code:
                flush_para()
                in_code = True
                code_lang = stripped[3:].strip()
                code_lines = []
            else:
                # 关闭代码块
                if code_lang.lower() not in ("mermaid",):  # 跳过 mermaid
                    add_code_block(doc, code_lines)
                in_code = False
                code_lang = ""
                code_lines = []
            continue

        if in_code:
            code_lines.append(line)
            continue

        # ---------- 表格行 ----------
        if stripped.startswith("|"):
            flush_para()
            in_table = True
            table_rows.append(stripped)
            continue
        else:
            if in_table:
                flush_table()

        # ---------- 空行 ----------
        if not stripped:
            flush_para()
            continue

        # ---------- 分隔线 ----------
        if stripped == "---":
            flush_para()
            doc.add_paragraph("─" * 40)
            continue

        # ---------- 标题 ----------
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            flush_para()
            level = min(len(m.group(1)), 4)
            doc.add_heading(clean_inline(m.group(2)), level=level)
            continue

        # ---------- 有序列表 ----------
        m = re.match(r"^\d+\.\s+(.+)$", stripped)
        if m:
            flush_para()
            doc.add_paragraph(clean_inline(m.group(1)), style="List Number")
            continue

        # ---------- 无序列表 ----------
        m = re.match(r"^[-*]\s+(.+)$", stripped)
        if m:
            flush_para()
            doc.add_paragraph(clean_inline(m.group(1)), style="List Bullet")
            continue

        # ---------- 普通段落 ----------
        para_buffer.append(stripped)

    # 收尾
    flush_para()
    if in_table:
        flush_table()

    doc.save(output_path)
    print(f"已生成：{output_path}")


def main() -> None:
    EXPORT_DIR.mkdir(exist_ok=True)
    text = SOURCE_MD.read_text(encoding="utf-8")
    export_docx(text, OUTPUT_DOCX)


if __name__ == "__main__":
    main()
